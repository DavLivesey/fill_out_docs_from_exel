import re
import openpyxl
import openpyxl.styles
import docx
import datetime as dt


path = "/path/to/directory/"
doc = docx.Document(f"{path}pattern.docx")


file = f"{path}source.xlsx"
font = openpyxl.styles.Font(size=14, name="TimesNewRoman")
  

today = dt.datetime.strftime(dt.date.today(), '%d-%m-%Y')

def extract_full_name(text):
    # Удаляем пробелы в начале и конце строки
    text = text.strip()
    # Ищем последнее вхождение имени и отчества
    match = re.search(r'([А-Я][а-я]+ [А-Я][а-я]+)', text)
    if match:
        # Возвращаем найденное значение, разворачивая его обратно
        return match.group(1)
    return None

def add_data():
        exel = openpyxl.load_workbook(file)
        list_1 = exel[exel.sheetnames[1]]
        list_1_count = 0
        list_2_count = 0
        list_3_count = 0
        list_2 = exel[exel.sheetnames[2]]
        list_3 = exel[exel.sheetnames[3]]
        for number_row in range(3, len(list_1['B'])):
                fullname = str(list_1.cell(row=number_row, column=2).value)
                result = extract_full_name(fullname)
                if result:
                        sex = str(list_1.cell(row=number_row, column=4).value)
                        doc.paragraphs[4].clear()
                        doc.paragraphs[4].style.font.name = 'Times New Roman'
                        doc.paragraphs[4].style.font.size = docx.shared.Pt(12)
                        if sex == "м":
                                doc.paragraphs[4].text = f"Глубокоуважаемый {result}!"
                        elif sex == "ж":
                              doc.paragraphs[4].text = f"Глубокоуважаемая {result}!"
                        else:
                              pass
                        doc.save(f"{path}{fullname}.docx")
                        list_1_count += 1
        for number_row in range(3, len(list_2['A'])):
                fullname = str(list_2.cell(row=number_row, column=1).value)
                result = (" ").join(fullname.strip().split(" ")[1::])
                if result:
                        sex = str(list_2.cell(row=number_row, column=3).value)
                        doc.paragraphs[4].clear()
                        doc.paragraphs[4].style.font.name = 'Times New Roman'
                        doc.paragraphs[4].style.font.size = docx.shared.Pt(12)
                        if sex == "м":
                                doc.paragraphs[4].text = f"Глубокоуважаемый {result}!"
                        elif sex == "ж":
                              doc.paragraphs[4].text = f"Глубокоуважаемая {result}!"
                        else:
                              pass
                        doc.save(f"{path}{fullname}.docx")
                        list_2_count += 1
        for number_row in range(3, len(list_3['A'])):
                fullname = str(list_3.cell(row=number_row, column=1).value)
                result = (" ").join(fullname.strip().split(" ")[1::])
                if result:
                        sex = str(list_3.cell(row=number_row, column=4).value)
                        doc.paragraphs[4].clear()
                        doc.paragraphs[4].style.font.name = 'Times New Roman'
                        doc.paragraphs[4].style.font.size = docx.shared.Pt(12)
                        if sex == "м":
                                doc.paragraphs[4].text = f"Глубокоуважаемый {result}!"
                        elif sex == "ж":
                              doc.paragraphs[4].text = f"Глубокоуважаемая {result}!"
                        else:
                              pass
                        doc.save(f"{path}{fullname}.docx")
                        list_3_count += 1
        print(list_1_count)
        print(list_2_count)
        print(list_3_count)
        


if __name__ == "__main__":
    add_data()
